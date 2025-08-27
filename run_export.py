import os, sys, csv, json, zipfile, time, uuid, math, argparse
from datetime import datetime, timezone, timedelta
import requests
from dateutil import parser as dtparser
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv

# ---------- config / env ----------
load_dotenv()
CLIENT_ID = os.getenv("CLIENT_ID", "").strip()
CLIENT_SECRET = os.getenv("CLIENT_SECRET", "").strip()
USER_AGENT = os.getenv("USER_AGENT", "linux:joci-reddit-scraper:1.0 (by /u/unknown)")

API = "https://oauth.reddit.com"
AUTH = "https://www.reddit.com/api/v1/access_token"
BASE_HEADERS = {"User-Agent": USER_AGENT, "Accept": "application/json"}

def ensure_ok():
    if not CLIENT_ID:
        sys.exit("Missing CLIENT_ID in .env")
    if not USER_AGENT or "by /u/" not in USER_AGENT:
        print("[warn] Your USER_AGENT should include your Reddit username per Reddit policy.")

def to_unix(iso_str):
    if iso_str.lower() == "now":
        return int(datetime.now(timezone.utc).timestamp())
    return int(dtparser.parse(iso_str).replace(tzinfo=timezone.utc).timestamp())

def get_token():
    # If CLIENT_SECRET present -> client_credentials; else installed_client
    if CLIENT_SECRET:
        auth = (CLIENT_ID, CLIENT_SECRET)
        data = {"grant_type": "client_credentials", "duration": "permanent"}
    else:
        auth = (CLIENT_ID, "")
        data = {"grant_type": "https://oauth.reddit.com/grants/installed_client",
                "device_id": str(uuid.uuid4())}
    r = requests.post(AUTH, data=data, headers=BASE_HEADERS, auth=auth, timeout=30)
    if r.status_code >= 400:
        raise SystemExit(f"Token error: {r.status_code} {r.text[:300]}")
    return r.json()["access_token"]

def h(tok):
    d = dict(BASE_HEADERS)
    d["Authorization"] = f"Bearer {tok}"
    return d

def sleep_reset(resp):
    if resp.status_code in (403, 429) or resp.headers.get("x-ratelimit-remaining") in ("0","0.0"):
        reset = resp.headers.get("x-ratelimit-reset")
        if reset:
            try:
                t = math.ceil(float(reset)) + 1
                print(f"[rate-limit] sleeping {t}s")
                time.sleep(max(5, t))
                return True
            except Exception:
                pass
        print("[rate-limit] sleeping 10s")
        time.sleep(10)
        return True
    return False

def month_windows(start_unix, end_unix):
    dt = datetime.fromtimestamp(start_unix, tz=timezone.utc).replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    end_dt = datetime.fromtimestamp(end_unix, tz=timezone.utc)
    while dt <= end_dt:
        nxt = (dt.replace(day=28) + timedelta(days=4)).replace(day=1)
        yield int(dt.timestamp()), int((nxt - timedelta(seconds=1)).timestamp())
        dt = nxt

def next_after(data):
    return data.get("data", {}).get("after")

def search_listing(token, sr, q, start_u, end_u, type_, limit=100, max_pages=10000):
    url = f"{API}/search"
    after = None
    got = []
    pages = 0
    while True:
        parts = []
        if q: parts.append(f"({q})")
        parts.append(f"timestamp:{start_u}..{end_u}")
        query = " AND ".join(parts)
        params = {
            "q": query, "restrict_sr": "on", "include_over_18": "on",
            "sort": "new", "syntax": "lucene", "type": type_,
            "limit": str(limit), "sr": sr
        }
        if after: params["after"] = after
        r = requests.get(url, headers=h(token), params=params, timeout=30)
        if r.status_code in (403, 429):
            if sleep_reset(r): continue
        r.raise_for_status()
        data = r.json()
        children = data.get("data", {}).get("children", [])
        got.extend(children)
        after = next_after(data)
        pages += 1
        if not after or not children or pages >= max_pages: break
    return got

def flatten_comments(children):
    out = []
    for ch in children:
        if ch.get("kind") != "t1": continue
        d = ch.get("data", {})
        out.append(d)
        repl = d.get("replies")
        if isinstance(repl, dict):
            out.extend(flatten_comments(repl.get("data", {}).get("children", [])))
    return out

def fetch_post_comments(token, sr, post_id):
    url = f"{API}/r/{sr}/comments/{post_id}.json"
    params = {"limit": "500"}
    r = requests.get(url, headers=h(token), params=params, timeout=30)
    if r.status_code in (403, 429):
        if sleep_reset(r): return fetch_post_comments(token, sr, post_id)
    r.raise_for_status()
    arr = r.json()
    if not isinstance(arr, list) or len(arr) < 2: return []
    return flatten_comments(arr[1].get("data", {}).get("children", []))

def collect(sr, query, from_iso, to_iso):
    ensure_ok()
    tok = get_token()
    start_u, end_u = to_unix(from_iso), to_unix(to_iso)
    comments = []
    # 1) Try direct comment search by month
    for su, eu in month_windows(start_u, end_u):
        for ch in search_listing(tok, sr, query, su, eu, type_="comment"):
            if ch.get("kind") == "t1":
                comments.append(ch["data"])
    # 2) Fallback: list posts then fetch each comment tree
    if not comments:
        posts = []
        for su, eu in month_windows(start_u, end_u):
            for ch in search_listing(tok, sr, query, su, eu, type_="link"):
                if ch.get("kind") == "t3":
                    posts.append(ch["data"])
        ids = list({p.get("id") for p in posts if p.get("id")})
        for pid in ids:
            comments.extend(fetch_post_comments(tok, sr, pid))
    # Dedup + filter
    seen, uniq = set(), []
    for c in comments:
        cid = c.get("id")
        if cid and cid not in seen:
            seen.add(cid)
            uniq.append(c)
    out = []
    for c in uniq:
        cu = c.get("created_utc")
        try:
            if cu is None or (to_unix(from_iso) <= int(cu) <= to_unix(to_iso)):
                out.append(c)
        except Exception:
            out.append(c)
    out.sort(key=lambda d: d.get("created_utc", 0))
    return out

def to_row(c):
    author = c.get("author") or "unknown"
    created = ""
    if c.get("created_utc") is not None:
        created = datetime.fromtimestamp(int(c["created_utc"]), tz=timezone.utc).isoformat().replace("+00:00","Z")
    permalink = c.get("permalink") or ""
    if permalink.startswith("/"):
        permalink = "https://www.reddit.com" + permalink
    return [c.get("id",""), author, created, permalink, c.get("parent_id",""), c.get("link_id",""), c.get("score",""), c.get("body","")]

def write_outputs(sub, query, comments, from_iso, to_iso):
    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    os.makedirs("exports", exist_ok=True)
    base = f"exports/reddit_comments_{sub}_{ts}"
    csv_path = base + ".csv"
    with open(csv_path, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f); w.writerow(["id","author","created_at","permalink","parent_id","link_id","score","body"])
        for c in comments: w.writerow(to_row(c))

    docx_path = base + ".docx"
    doc = Document()
    doc.add_heading(f"Reddit comments for r/{sub}", level=0)
    doc.add_paragraph(f"Window: {from_iso} → {to_iso}")
    if query: doc.add_paragraph(f"Query: {query}")
    for c in comments:
        cid, author, created, permalink, parent, link, score, body = to_row(c)
        doc.add_heading(f"{author} • {created or 'unknown'} • score {score}", level=3)
        meta = doc.add_paragraph(); meta.add_run(f"URL: {permalink}\nID: {cid}\nParent: {parent}\nLink: {link}").italic = True
        p = doc.add_paragraph(body)
        for run in p.runs: run.font.name = "Courier New"; run.font.size = Pt(10)
        doc.add_paragraph().add_run("—").bold = True
    doc.save(docx_path)

    meta_path = base + "_meta.json"
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump({"subreddit": sub, "query": query, "from": from_iso, "to": to_iso, "counts": {"total": len(comments)}}, f, ensure_ascii=False, indent=2)

    readme_txt = base + "_README.txt"
    with open(readme_txt, "w", encoding="utf-8") as f:
        f.write(f"Reddit export for r/{sub}\nWindow: {from_iso} -> {to_iso}\nTotal comments: {len(comments)}\n")

    zip_path = base + ".zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        z.write(csv_path,    arcname=os.path.basename(csv_path))
        z.write(docx_path,   arcname=os.path.basename(docx_path))
        z.write(meta_path,   arcname=os.path.basename(meta_path))
        z.write(readme_txt,  arcname=os.path.basename(readme_txt))
    return zip_path

def main():
    ap = argparse.ArgumentParser(description="Reddit comment exporter (ZIP)")
    ap.add_argument("--subreddit", required=True, help="Subreddit name without r/")
    ap.add_argument("--query", default="", help="Optional text filter")
    ap.add_argument("--from", dest="from_iso", default="2012-01-01", help="Start date (ISO)")
    ap.add_argument("--to", dest="to_iso", default="now", help="End date (ISO or 'now')")
    args = ap.parse_args()

    comments = collect(args.subreddit, args.query, args.from_iso, args.to_iso)
    zip_path = write_outputs(args.subreddit, args.query, comments, args.from_iso, args.to_iso)
    print(f"[DONE] ZIP: {zip_path}")

if __name__ == "__main__":
    main()
